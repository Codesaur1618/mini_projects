import tkinter as tk
from tkinter import filedialog
from docx import Document
from fontTools.ttLib import TTFont

def extract_text_from_docx(docx_path):
    try:
        text = docx2txt.process(docx_path)
        return text
    except Exception as e:
        return str(e)

def find_fonts_used(docx_path):
    extracted_text = extract_text_from_docx(docx_path)

    if extracted_text:
        fonts_used = set()
        for paragraph in extracted_text.split('\n'):
            # Create a list of unique fonts used in each paragraph
            fonts = set()
            doc = Document()
            run = doc.add_paragraph(paragraph).runs[0]
            font_name = run.font.name
            fonts.add(font_name)

            # Check for additional fonts in the paragraph
            for run in doc.paragraphs[0].runs:
                font_name = run.font.name
                fonts.add(font_name)

            # Add the fonts used in the paragraph to the overall set
            fonts_used.update(fonts)

        return fonts_used
    else:
        return None

def open_file_dialog():
    root = tk.Tk()
    root.withdraw()  # Hide the main tkinter window

    file_path = filedialog.askopenfilename(
        title="Select a Document or Text File",
        filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx"), ("All Files", "*.*")]
    )

    if file_path:
        print("Selected file:", file_path)

        if file_path.lower().endswith('.docx'):
            fonts_used = find_fonts_used(file_path)
            if fonts_used:
                print("Fonts used in the document:")
                for font in fonts_used:
                    print(font)
            else:
                print("Failed to extract text from the document.")
        else:
            print("This program currently supports only DOCX files.")

if __name__ == "__main__":
    open_file_dialog()
